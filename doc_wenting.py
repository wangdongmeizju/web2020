from docx import Document
import time,datetime,pytz
from docxcompose.composer import Composer



def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.font.size=run.font.size
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

input_doc = Document('template.docx')


# get_para_data(output_doc, input_doc.paragraphs[2])

def hebing(files,final_docx):
    new_document = Document()
    composer = Composer(new_document)
    for index, sub_doc in enumerate(files):
        if index < len(files)-1:
           sub_doc.add_page_break()
        composer.append(sub_doc)
    composer.save(final_docx)

def process(txt,download_file_list,tijiaoren):
    time_wang=datetime.datetime.fromtimestamp(int(time.time()), pytz.timezone('Asia/Shanghai')).strftime('%Y-%m-%d %H:%M:%S')
    file_name=download_file_list+"/"+tijiaoren+time_wang+"banzheng.doc"
    infos=txt.split("@@@")
    kk=1
    output_docs=[]
    for info in infos:
        output_doc = Document()
        temp= info.split("；")
        if (len(temp)<4):
            return "1 校区,日期,车牌,人名是不是四个中缺了呀 查查看呀    2 ；是中文下的 ",""
        xiaoqu=temp[0]
        riqi=temp[1]
        chepai=temp[2]
        persons=temp[3]
        j=0
        i=0
        while i < len(input_doc.paragraphs):
            if i==2:
                for person in persons.split("，"):
                    if person != "":
                        get_para_data(output_doc, input_doc.paragraphs[2])
                        get_para_data(output_doc, input_doc.paragraphs[3])
                i=i+2
            get_para_data(output_doc, input_doc.paragraphs[i])
            i=i+1
        paragraphs=output_doc.paragraphs
        content_wen=[]
        for person in persons.split("，"):
            if person != "":
                temp=person.split("：")
                if (len(temp)!=2):
                    return "1 每个人都要身份证 人名 两样是否全了   2  ：是中文下的,",""
                p_n=temp[0]
                p_s=temp[1]
                content_wen.append(p_n)
                content_wen.append(p_s)
        content_wen.append(xiaoqu)
        content_wen.append(riqi)
        content_wen.append(chepai)

        i=0
        for ele in content_wen:
            paragraphs[i+2].add_run(content_wen[i])
            i=i+1

        output_docs.append(output_doc)
        kk=kk+1
    hebing(output_docs,file_name)

    return "success",file_name